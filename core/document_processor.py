"""
Document Processing Service for OAI_BOT_V.1
Handles PDF, DOC, DOCX, and TXT file processing
"""

import os
import logging
import uuid
from typing import List, Dict, Any, Optional
from datetime import datetime
import re

# PDF Processing
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

# DOC/DOCX Processing  
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Image Processing for OCR
try:
    from PIL import Image
    PILLOW_AVAILABLE = True
except ImportError:
    PILLOW_AVAILABLE = False

logger = logging.getLogger(__name__)

class DocumentProcessor:
    def __init__(self, settings):
        self.settings = settings
        self.supported_formats = {
            '.pdf': self._process_pdf,
            '.doc': self._process_doc,
            '.docx': self._process_docx,
            '.txt': self._process_txt
        }
        
        # Ensure upload directory exists
        os.makedirs(settings.upload_folder, exist_ok=True)
    
    async def process_document(self, file_path: str, filename: str) -> Dict[str, Any]:
        """Process document and extract text"""
        try:
            # Get file extension
            _, ext = os.path.splitext(filename.lower())
            
            if ext not in self.supported_formats:
                return {
                    "success": False,
                    "error": f"Unsupported file format: {ext}",
                    "supported_formats": list(self.supported_formats.keys())
                }
            
            # Generate document ID
            document_id = str(uuid.uuid4())
            
            logger.info(f"Processing document: {filename} (ID: {document_id})")
            
            # Process based on file type
            processor = self.supported_formats[ext]
            text_content = await processor(file_path)
            
            if not text_content:
                return {
                    "success": False,
                    "error": "No text content extracted from document"
                }
            
            # Split text into chunks
            chunks = self._split_text_into_chunks(text_content)
            
            # Prepare metadata
            metadata = {
                "document_id": document_id,
                "filename": filename,
                "file_size": os.path.getsize(file_path),
                "processed_at": datetime.now().isoformat(),
                "total_chunks": len(chunks),
                "total_characters": len(text_content)
            }
            
            logger.info(f"Document processed: {len(chunks)} chunks, {len(text_content)} characters")
            
            return {
                "success": True,
                "document_id": document_id,
                "filename": filename,
                "text_content": text_content,
                "chunks": chunks,
                "metadata": metadata
            }
            
        except Exception as e:
            logger.error(f"Error processing document {filename}: {str(e)}")
            return {
                "success": False,
                "error": f"Processing error: {str(e)}"
            }
    
    async def _process_pdf(self, file_path: str) -> str:
        """Process PDF file"""
        if not PDF_AVAILABLE:
            raise Exception("PyPDF2 not available. Install with: pip install PyPDF2")
        
        try:
            text_content = ""
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text_content += f"\n--- หน้า {page_num + 1} ---\n"
                            text_content += page_text
                    except Exception as e:
                        logger.warning(f"Error extracting text from page {page_num + 1}: {str(e)}")
                        continue
            
            return self._clean_text(text_content)
            
        except Exception as e:
            logger.error(f"Error processing PDF: {str(e)}")
            raise
    
    async def _process_docx(self, file_path: str) -> str:
        """Process DOCX file"""
        if not DOCX_AVAILABLE:
            raise Exception("python-docx not available. Install with: pip install python-docx")
        
        try:
            doc = Document(file_path)
            text_content = ""
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content += paragraph.text + "\n"
            
            # Process tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content += " | ".join(row_text) + "\n"
            
            return self._clean_text(text_content)
            
        except Exception as e:
            logger.error(f"Error processing DOCX: {str(e)}")
            raise
    
    async def _process_doc(self, file_path: str) -> str:
        """Process DOC file (older Word format)"""
        # For DOC files, we'll need additional libraries or conversion
        # For now, return an error message
        raise Exception("DOC format not yet supported. Please convert to DOCX or PDF.")
    
    async def _process_txt(self, file_path: str) -> str:
        """Process TXT file"""
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-8-sig', 'tis-620', 'cp874']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        content = file.read()
                    logger.info(f"Successfully read TXT file with {encoding} encoding")
                    return self._clean_text(content)
                except UnicodeDecodeError:
                    continue
            
            # If all encodings fail, try with error handling
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                content = file.read()
                logger.warning("Read TXT file with ignored encoding errors")
                return self._clean_text(content)
            
        except Exception as e:
            logger.error(f"Error processing TXT: {str(e)}")
            raise
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize text"""
        if not text:
            return ""
        
        # Remove excessive whitespace
        text = re.sub(r'\n\s*\n', '\n\n', text)
        text = re.sub(r'[ \t]+', ' ', text)
        
        # Remove special characters that might cause issues
        text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x84\x86-\x9f]', '', text)
        
        # Ensure proper line breaks
        text = text.replace('\r\n', '\n').replace('\r', '\n')
        
        return text.strip()
    
    def _split_text_into_chunks(self, text: str) -> List[Dict[str, Any]]:
        """Split text into smaller chunks for embedding"""
        if not text:
            return []
        
        chunk_size = self.settings.chunk_size
        overlap = self.settings.chunk_overlap
        
    def _split_text_into_chunks(self, text: str) -> List[Dict[str, Any]]:
        """Split text into smaller chunks for embedding"""
        if not text:
            return []
        
        chunk_size = self.settings.chunk_size
        overlap = self.settings.chunk_overlap
        
        # Split by paragraphs first
        paragraphs = text.split('\n\n')
        chunks = []
        current_chunk = ""
        chunk_index = 0
        
        for paragraph in paragraphs:
            paragraph = paragraph.strip()
            if not paragraph:
                continue
            
            # If adding this paragraph exceeds chunk size, save current chunk
            if len(current_chunk) + len(paragraph) > chunk_size and current_chunk:
                chunks.append({
                    "text": current_chunk.strip(),
                    "chunk_index": chunk_index,
                    "character_count": len(current_chunk.strip())
                })
                
                # Start new chunk with overlap
                if overlap > 0 and len(current_chunk) > overlap:
                    current_chunk = current_chunk[-overlap:] + "\n\n" + paragraph
                else:
                    current_chunk = paragraph
                
                chunk_index += 1
            else:
                # Add paragraph to current chunk
                if current_chunk:
                    current_chunk += "\n\n" + paragraph
                else:
                    current_chunk = paragraph
        
        # Add the last chunk
        if current_chunk.strip():
            chunks.append({
                "text": current_chunk.strip(),
                "chunk_index": chunk_index,
                "character_count": len(current_chunk.strip())
            })
        
        # If no chunks were created (text too short), create one chunk
        if not chunks and text.strip():
            chunks.append({
                "text": text.strip(),
                "chunk_index": 0,
                "character_count": len(text.strip())
            })
        
        logger.info(f"Text split into {len(chunks)} chunks")
        return chunks
    
    def save_uploaded_file(self, file_data: bytes, filename: str) -> str:
        """Save uploaded file to disk"""
        try:
            # Generate unique filename to avoid conflicts
            base_name, ext = os.path.splitext(filename)
            unique_filename = f"{base_name}_{uuid.uuid4().hex[:8]}{ext}"
            file_path = os.path.join(self.settings.upload_folder, unique_filename)
            
            with open(file_path, 'wb') as f:
                f.write(file_data)
            
            logger.info(f"File saved: {file_path}")
            return file_path
            
        except Exception as e:
            logger.error(f"Error saving file {filename}: {str(e)}")
            raise
    
    def delete_file(self, file_path: str) -> bool:
        """Delete file from disk"""
        try:
            if os.path.exists(file_path):
                os.remove(file_path)
                logger.info(f"File deleted: {file_path}")
                return True
            return False
        except Exception as e:
            logger.error(f"Error deleting file {file_path}: {str(e)}")
            return False
    
    def get_file_info(self, file_path: str) -> Dict[str, Any]:
        """Get file information"""
        try:
            if not os.path.exists(file_path):
                return {"exists": False}
            
            stat = os.stat(file_path)
            return {
                "exists": True,
                "size": stat.st_size,
                "created": datetime.fromtimestamp(stat.st_ctime).isoformat(),
                "modified": datetime.fromtimestamp(stat.st_mtime).isoformat()
            }
        except Exception as e:
            logger.error(f"Error getting file info: {str(e)}")
            return {"exists": False, "error": str(e)}
    
    def validate_file(self, filename: str, file_size: int) -> Dict[str, Any]:
        """Validate uploaded file"""
        errors = []
        
        # Check file extension
        if not self.settings.is_allowed_file(filename):
            errors.append(f"File type not allowed. Supported: {', '.join(self.settings.allowed_extensions)}")
        
        # Check file size
        if file_size > self.settings.max_file_size:
            max_mb = self.settings.max_file_size / (1024 * 1024)
            errors.append(f"File too large. Maximum size: {max_mb:.1f} MB")
        
        # Check filename
        if not filename or filename.strip() == "":
            errors.append("Filename cannot be empty")
        
        # Check for dangerous characters
        dangerous_chars = ['..', '/', '\\', '<', '>', ':', '"', '|', '?', '*']
        if any(char in filename for char in dangerous_chars):
            errors.append("Filename contains invalid characters")
        
        return {
            "valid": len(errors) == 0,
            "errors": errors
        }
    
    def health_check(self) -> bool:
        """Check if document processor is healthy"""
        try:
            # Check if upload directory exists and is writable
            if not os.path.exists(self.settings.upload_folder):
                return False
            
            # Try to create a test file
            test_file = os.path.join(self.settings.upload_folder, "health_check.tmp")
            with open(test_file, 'w') as f:
                f.write("test")
            
            # Clean up test file
            os.remove(test_file)
            
            return True
            
        except Exception as e:
            logger.error(f"Document processor health check failed: {str(e)}")
            return False
    
    async def extract_metadata(self, file_path: str, filename: str) -> Dict[str, Any]:
        """Extract metadata from document"""
        try:
            metadata = {
                "filename": filename,
                "file_size": os.path.getsize(file_path),
                "created_at": datetime.now().isoformat()
            }
            
            # Get file extension
            _, ext = os.path.splitext(filename.lower())
            metadata["file_type"] = ext
            
            # Extract specific metadata based on file type
            if ext == '.pdf' and PDF_AVAILABLE:
                try:
                    with open(file_path, 'rb') as file:
                        pdf_reader = PyPDF2.PdfReader(file)
                        metadata["page_count"] = len(pdf_reader.pages)
                        
                        # Try to get PDF metadata
                        if pdf_reader.metadata:
                            metadata["pdf_title"] = pdf_reader.metadata.get("/Title", "")
                            metadata["pdf_author"] = pdf_reader.metadata.get("/Author", "")
                            metadata["pdf_subject"] = pdf_reader.metadata.get("/Subject", "")
                except Exception as e:
                    logger.warning(f"Could not extract PDF metadata: {str(e)}")
            
            elif ext == '.docx' and DOCX_AVAILABLE:
                try:
                    doc = Document(file_path)
                    metadata["paragraph_count"] = len(doc.paragraphs)
                    metadata["table_count"] = len(doc.tables)
                    
                    # Try to get document properties
                    if doc.core_properties:
                        metadata["docx_title"] = doc.core_properties.title or ""
                        metadata["docx_author"] = doc.core_properties.author or ""
                        metadata["docx_subject"] = doc.core_properties.subject or ""
                except Exception as e:
                    logger.warning(f"Could not extract DOCX metadata: {str(e)}")
            
            return metadata
            
        except Exception as e:
            logger.error(f"Error extracting metadata: {str(e)}")
            return {
                "filename": filename,
                "error": str(e)
            }
    
    def get_supported_formats(self) -> Dict[str, Dict[str, Any]]:
        """Get information about supported file formats"""
        formats = {}
        
        for ext, processor in self.supported_formats.items():
            available = True
            requirements = []
            
            if ext == '.pdf':
                available = PDF_AVAILABLE
                requirements = ["PyPDF2"]
            elif ext in ['.doc', '.docx']:
                available = DOCX_AVAILABLE if ext == '.docx' else False
                requirements = ["python-docx"] if ext == '.docx' else ["Additional library required"]
            
            formats[ext] = {
                "available": available,
                "requirements": requirements,
                "description": self._get_format_description(ext)
            }
        
        return formats
    
    def _get_format_description(self, ext: str) -> str:
        """Get description for file format"""
        descriptions = {
            '.pdf': "Portable Document Format - รองรับการแยกข้อความและตาราง",
            '.docx': "Microsoft Word Document (2007+) - รองรับข้อความ ตาราง และรูปแบบ",
            '.doc': "Microsoft Word Document (Legacy) - ต้องการการแปลงไฟล์",
            '.txt': "Plain Text File - รองรับการเข้ารหัสภาษาไทย"
        }
        return descriptions.get(ext, "Unknown format")
